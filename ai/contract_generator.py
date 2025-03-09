from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import base64
import io
from PIL import Image, ImageDraw, ImageFont
from openai import OpenAI
from django.conf import settings
from datetime import datetime


def generate_contract(data):
    """
    Generate a professional employment contract using AI in markdown format,
    then convert it to a formatted Word document with pre-filled employer signature

    Args:
        data (dict): Dictionary containing all contract details

    Returns:
        str: Path to the generated contract file
    """
    # Create a prompt that explicitly requests markdown formatting
    prompt = f"""
    Generate a professional employment contract in markdown format with the following details:

    EMPLOYER INFORMATION:
    - Employer: {data['employer_name']}
    - Principal office: {data['employer_address']}

    EMPLOYEE INFORMATION:
    - Employee: {data['employee_name']}
    - Employee address: {data['employee_address']}

    EMPLOYMENT DETAILS:
    - Job Title: {data['job_title']}
    - Start Date: {data['start_date']}
    - End Date: {data.get('end_date', 'Not specified - employment at will')}
    - Salary: {data['salary']}
    
    JOB DETAILS:
    - Responsibilities: {data['responsibilities']}
    - Benefits: {data.get('benefits', 'Standard benefits package as per company policy')}
    
    ADDITIONAL TERMS:
    {data['terms']}

    The contract should be formal, professional, and legally structured, following these sections:
    1. Introduction
    2. Parties and Employment Position
    3. Term of Employment
    4. Compensation and Benefits
    5. Duties and Responsibilities
    6. Confidentiality and Intellectual Property
    7. Termination
    8. Governing Law
    9. Entire Agreement
    10. Signatures

    Use proper markdown formatting:
    - # for the main title
    - ## for major section headers (like "Article I: Parties")
    - ### for subsection headers
    - **bold** for emphasis
    - --- for horizontal dividers between sections
    - Proper numbered lists for sections like 1.1, 1.2, etc.
    - Ensure proper structure and layout in the markdown
    """

    try:
        client = OpenAI(
            base_url=settings.OPENAI_ENDPOINT,
            api_key=settings.OPENAI_TOKEN,
        )

        completion = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You are a legal expert specialized in employment contracts. Generate a complete employment contract in markdown format with proper structure and formatting.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )

        # Get the markdown content
        markdown_content = completion.choices[0].message.content.strip()

        # Remove code block markers if present
        if markdown_content.startswith("```markdown"):
            markdown_content = "\n".join(markdown_content.split("\n")[1:])
        if markdown_content.startswith("```"):
            markdown_content = "\n".join(markdown_content.split("\n")[1:])
        if markdown_content.endswith("```"):
            markdown_content = "\n".join(markdown_content.split("\n")[:-1])
            
        # Ensure the contracts directory exists
        contracts_dir = os.path.join(settings.MEDIA_ROOT, "contracts")
        os.makedirs(contracts_dir, exist_ok=True)

        # Create a unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_employee_name = "".join(
            c if c.isalnum() else "_" for c in data["employee_name"]
        )
        
        # Create the Word document from markdown with pre-filled employer signature
        document = markdown_to_docx(markdown_content, data)
        
        contract_filename = f"{safe_employee_name}_{timestamp}_contract.docx"
        contract_path = os.path.join(contracts_dir, contract_filename)

        document.save(contract_path)
        return contract_path

    except Exception as e:
        print(f"Contract generation error: {str(e)}")
        raise


def markdown_to_docx(markdown_text, data):
    """
    Convert markdown text to a formatted Word document
    
    Args:
        markdown_text (str): The markdown text to convert
        data (dict): Contract data for custom fields
        
    Returns:
        Document: A python-docx Document object
    """
    # Create document with basic styles
    document = Document()
    setup_document_styles(document)
    
    # Set margins
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    
    # Add header with the current date
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}").font.size = Pt(9)
    
    # Add footer with page numbers (fixed implementation)
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(footer_para)
    
    # Split the markdown content into lines
    lines = markdown_text.split('\n')
    i = 0
    
    # Process the markdown line by line
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
            
        # Process headings
        if line.startswith('# '):
            # Main title (h1)
            text = line[2:].strip()
            heading = document.add_heading(text, 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
        elif line.startswith('## '):
            # Article heading (h2)
            text = line[3:].strip()
            heading = document.add_heading(text, 1)
            i += 1
        elif line.startswith('### '):
            # Sub-section heading (h3)
            text = line[4:].strip()
            heading = document.add_heading(text, 2)
            i += 1
        elif line.startswith('---'):
            # Horizontal line
            p = document.add_paragraph()
            p.paragraph_format.border_bottom = True
            p.paragraph_format.space_after = Pt(12)
            i += 1
        elif re.match(r'^\d+\.\d+', line):
            # Section number (like 1.1)
            p = document.add_paragraph(style='ListNumber')
            process_markdown_text(p, line)
            i += 1
        else:
            # Regular paragraph
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            process_markdown_text(p, line)
            i += 1
            
            # Check if this is part of a multi-line paragraph
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '-', '*', '1.', '---')):
                process_markdown_text(p, ' ' + lines[i].strip())
                i += 1
    
    # Add signature section with pre-filled employer signature
    add_signature_section(document, data)
    
    return document


def process_markdown_text(paragraph, text):
    """Process markdown text with bold formatting"""
    # Handle bold text: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        bold_match = re.match(r'\*\*(.*?)\*\*$', part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            run.font.size = Pt(11)
        elif part:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def setup_document_styles(document):
    """Set up document styles for headings and paragraphs"""
    styles = document.styles
    
    # Style for title (Heading 0)
    title_style = styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
    
    # Style for heading 1
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
    
    # Style for heading 2
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    
    # Style for normal text
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    
    # Create a list style
    if 'ListNumber' not in styles:
        list_style = styles.add_style('ListNumber', WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.name = 'Times New Roman'
        list_style.font.size = Pt(11)
        list_style.font.bold = True
        
    # Create table style
    table_style = document.styles.add_style('TableGrid', WD_STYLE_TYPE.TABLE)


def add_signature_section(document, data):
    """Add signature section to the document with employer signature pre-filled"""
    document.add_page_break()
    
    # Add signature heading
    heading = document.add_heading('SIGNATURES', 1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    document.add_paragraph()
    
    # Create a table for signatures - 2 rows, 2 columns
    table = document.add_table(rows=2, cols=2)
    table.style = 'Table Grid'  # Add borders to the table
    
    # Ensure the table takes up most of the page width
    table.autofit = False
    for column in table.columns:
        column.width = Inches(3.0)
    
    # Generate employer signature image
    employer_signature = generate_signature_image(data['employer_name'])
    
    # Current date
    current_date = datetime.now().strftime('%B %d, %Y')
    
    # Employer signature - already signed
    cell = table.cell(0, 0)
    cell.width = Inches(3.0)
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("For the Employer:\n").bold = True
    
    # Add the signature image
    try:
        run = p.add_run()
        run.add_picture(employer_signature, width=Inches(2.5))
        
        # Add employer name and date
        emp_run = p.add_run(f"\n{data['employer_name']}\n")
        emp_run.bold = True
        emp_run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
        
        date_run = p.add_run(f"Date: {current_date}")
        date_run.italic = True
    except Exception as e:
        # Fallback if image insertion fails
        print(f"Error adding signature image: {e}")
        p.add_run("\n[Signed]\n")
        p.add_run(f"{data['employer_name']}\n")
        p.add_run(f"Date: {current_date}")
    
    # Employee signature - needs to be signed
    cell = table.cell(0, 1)
    cell.width = Inches(3.0)
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("Employee:\n").bold = True
    p.add_run("_________________________\n")
    emp_name = p.add_run(data['employee_name'] + "\n")
    emp_name.bold = True
    p.add_run("Date: ___________________")
    
    # Witness signatures
    for col in range(2):
        cell = table.cell(1, col)
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("Witness:\n").bold = True
        p.add_run("_________________________\n")
        p.add_run("Name: ___________________\n")
        p.add_run("Date: ___________________")


def generate_signature_image(name):
    """
    Generate a signature-like image for the given name
    
    Args:
        name (str): Name to convert to signature
    
    Returns:
        BytesIO: Image file-like object containing the signature
    """
    # Create a white image with a solid white background (not transparent)
    width, height = 500, 150
    image = Image.new('RGB', (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(image)
    
    # Create a signature style text
    try:
        # Try various fallback fonts
        font = None
        font_size = 50
        
        # List of possible script/cursive fonts that might be available on the system
        possible_fonts = [
            os.path.join(settings.BASE_DIR, 'ai/fonts/signature.ttf'),
            "Segoe Script", "Lucida Handwriting", "Brush Script MT", 
            "Comic Sans MS", "Ink Free", "Mistral", "Bradley Hand ITC",
            "Freestyle Script"
        ]
        
        # Try each font until one works
        for font_name in possible_fonts:
            try:
                if font_name.endswith('.ttf') and os.path.exists(font_name):
                    font = ImageFont.truetype(font_name, font_size)
                    break
                else:
                    font = ImageFont.truetype(font_name, font_size)
                    break
            except Exception:
                continue
        
        # If no font worked, fallback to default
        if font is None:
            # Try default system font
            try:
                # Different ways to get a default font on various systems
                default_font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",  # Linux
                    "/System/Library/Fonts/Times.ttc",                    # macOS
                    "C:/Windows/Fonts/times.ttf",                         # Windows
                    "arial.ttf",                                          # Generic
                ]
                
                for font_path in default_font_paths:
                    if os.path.exists(font_path):
                        font = ImageFont.truetype(font_path, font_size)
                        break
            except Exception:
                # Last resort: use default
                font = ImageFont.load_default()
                font_size = 30  # Default font is smaller
        
        # Create a more signature-like appearance
        # First add a subtle light blue background under the signature area
        draw.rectangle([20, 40, width-20, 110], fill=(240, 248, 255))
        
        # Draw the signature text
        signature_text = name
        
        # If we have a proper font, make it look more signature-like
        if font is not None and font != ImageFont.load_default():
            # Use a dark blue color
            text_color = (0, 0, 139)  # Dark blue
            
            try:
                # Get text width for centering (handled for both old and new Pillow versions)
                try:
                    # For newer Pillow versions
                    text_width = draw.textlength(signature_text, font=font)
                except AttributeError:
                    # For older Pillow versions
                    text_width, _ = draw.textsize(signature_text, font=font)
                
                text_x = (width - text_width) // 2
            except Exception:
                # If text width calculation fails, use default position
                text_x = 50
                
            # Draw the signature with a slight slant for realism
            draw.text((text_x, 50), signature_text, font=font, fill=text_color)
            
            # Add a small line under the signature for professional look
            line_y = 105
            draw.line([(text_x, line_y), (text_x + min(350, text_width * 1.2), line_y)], 
                      fill=text_color, width=1)
        else:
            # Fallback: simple text with underline
            draw.text((50, 50), signature_text, fill=(0, 0, 139))
            draw.line([(50, 85), (350, 85)], fill=(0, 0, 139), width=1)
        
    except Exception as e:
        # If all else fails, create a very basic signature
        print(f"Error creating signature image: {e}")
        draw.text((50, 50), f"{name} (signed)", fill=(0, 0, 0))
        draw.line([(50, 85), (350, 85)], fill=(0, 0, 0), width=1)
    
    # Save the signature to a BytesIO object
    signature_io = io.BytesIO()
    image.save(signature_io, format='PNG')
    signature_io.seek(0)
    
    return signature_io


def add_page_number(paragraph):
    """Add page numbers in format 'Page X of Y' using proper XML elements"""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    
    # Add "Page "
    run.add_text("Page ")
    
    # Add page number field
    run = paragraph.add_run()
    run.font.size = Pt(9)
    
    # Create the field for PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    # Add " of "
    run = paragraph.add_run(" of ")
    run.font.size = Pt(9)
    
    # Create the field for NUMPAGES
    run = paragraph.add_run()
    run.font.size = Pt(9)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = " NUMPAGES "
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar3)
    run._element.append(instrText2)
    run._element.append(fldChar4)
